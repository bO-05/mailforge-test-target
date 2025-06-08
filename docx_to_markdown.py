```python
import docx
import re
import os

def convert_docx_to_markdown(docx_path, md_path, images_folder='images'):
    # Open the .docx file
    doc = docx.Document(docx_path)

    # Initialize the Markdown content
    md_content = []

    # Create images folder if it doesn't exist
    if not os.path.exists(images_folder):
        os.makedirs(images_folder)

    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        # Convert paragraph text to Markdown
        md_para = convert_paragraph_to_markdown(para, images_folder)
        md_content.append(md_para)

    # Convert tables to Markdown
    for table in doc.tables:
        md_table = convert_table_to_markdown(table)
        md_content.append(md_table)

    # Write the Markdown content to a .md file
    with open(md_path, 'w', encoding='utf-8') as md_file:
        md_file.write('\n'.join(md_content))

def convert_paragraph_to_markdown(paragraph, images_folder):
    # Initialize the Markdown paragraph
    md_para = paragraph.text

    # Convert headings
    if paragraph.style.name.startswith('Heading'):
        level = int(paragraph.style.name.split()[-1])
        md_para = f"{'#' * level} {md_para}"

    # Convert bold text
    for run in paragraph.runs:
        if run.bold:
            md_para = md_para.replace(run.text, f"**{run.text}**")

    # Convert italic text
    for run in paragraph.runs:
        if run.italic:
            md_para = md_para.replace(run.text, f"*{run.text}*")

    # Convert bullet points
    if paragraph.style.name == 'List Paragraph':
        md_para = f"- {md_para}"

    # Convert hyperlinks
    for run in paragraph.runs:
        if run.hyperlink:
            url = run.hyperlink.address
            md_para = md_para.replace(run.text, f"[{run.text}]({url})")

    # Convert images
    for run in paragraph.runs:
        if run.element.tag.endswith('drawing'):
            image_path = save_image(run, images_folder)
            md_para = md_para.replace(run.text, f"![{run.text}]({image_path})")

    # Convert code blocks
    if paragraph.style.name == 'Code' or any(run.font.name == 'Courier New' for run in paragraph.runs):
        md_para = f"```\n{md_para}\n```"

    return md_para

def convert_table_to_markdown(table):
    md_table = []
    for row in table.rows:
        md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
        md_table.append(md_row)
        if row == table.rows[0]:
            md_table.append('| ' + ' | '.join('---' for _ in row.cells) + ' |')
    return '\n'.join(md_table)

def save_image(run, images_folder):
    image = run.element.xpath('.//a:blipFill/a:blip/@r:embed', namespaces=run.element.nsmap)[0]
    image_part = run.part.related_parts[image]
    image_path = os.path.join(images_folder, f'image{len(os.listdir(images_folder)) + 1}.png')
    with open(image_path, 'wb') as img_file:
        img_file.write(image_part.blob)
    return image_path

# Example usage
if __name__ == "__main__":
    docx_path = 'example.docx'
    md_path = 'example.md'
    convert_docx_to_markdown(docx_path, md_path)
```